"""J.3 / J.4 Word modern + legacy."""
from __future__ import annotations

import asyncio
import io
import logging
import zipfile

from defusedxml.ElementTree import fromstring as et_fromstring
from docx import Document
from docx.oxml.ns import qn

from trimcp.extractors.common import is_zip_encrypted_ooxml
from trimcp.extractors.core import ExtractionResult, Section, empty_skipped
from trimcp.extractors.libreoffice import libreoffice_convert

log = logging.getLogger(__name__)


def _extract_core_props_xml(blob: bytes) -> dict:
    meta: dict = {}
    try:
        with zipfile.ZipFile(io.BytesIO(blob)) as z:
            if "docProps/core.xml" not in z.namelist():
                return meta
            root = et_fromstring(z.read("docProps/core.xml"))
            for el in root.iter():
                local = el.tag.split("}")[-1] if "}" in el.tag else el.tag
                if el.text and local in ("creator", "created", "modified", "lastModifiedBy"):
                    meta[local] = el.text.strip()
    except Exception as e:
        log.debug("core_props: %s", e)
    return meta


async def extract_docx(blob: bytes) -> ExtractionResult:
    return await asyncio.to_thread(extract_docx_sync, blob)


def extract_docx_sync(blob: bytes) -> ExtractionResult:
    if is_zip_encrypted_ooxml(blob):
        return empty_skipped("python-docx", "encrypted")
    warnings: list[str] = []
    try:
        doc = Document(io.BytesIO(blob))
    except Exception as e:
        log.warning("python-docx failed, trying LibreOffice: %s", e)
        converted = libreoffice_convert(blob, ".docx", ".docx")
        if not converted:
            return empty_skipped(
                "libreoffice",
                "conversion_failed",
                warnings=[f"python-docx failed: {e}", "libreoffice fallback failed"],
            )
        try:
            doc = Document(io.BytesIO(converted))
            warnings.append("Opened via LibreOffice re-normalization after python-docx failure")
        except Exception as e2:
            return empty_skipped("python-docx", "corrupt", warnings=[str(e), str(e2)])

    sections: list[Section] = []
    order = 0
    heading_stack: list[str] = []

    for para in doc.paragraphs:
        try:
            if not para.text.strip():
                continue
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    level_s = style_name.replace("Heading", "").strip()
                    level = int(level_s) if level_s else 1
                except ValueError:
                    level = 1
                heading_stack = heading_stack[: max(0, level - 1)] + [para.text.strip()]
                sections.append(
                    Section(
                        text=para.text,
                        structure_path=" > ".join(heading_stack),
                        section_type="heading",
                        order=order,
                    )
                )
            else:
                path = " > ".join(heading_stack) if heading_stack else "Body"
                sections.append(Section(text=para.text, structure_path=path, section_type="body", order=order))
            order += 1
        except Exception as e:
            warnings.append(f"paragraph_skip:{e}")

    for tbl in doc.tables:
        try:
            rows = ["| " + " | ".join(cell.text.replace("|", "\\|") for cell in row.cells) + " |" for row in tbl.rows]
            path = (" > ".join(heading_stack) + " > Table") if heading_stack else "Table"
            sections.append(Section(text="\n".join(rows), structure_path=path, section_type="table", order=order))
            order += 1
        except Exception as e:
            warnings.append(f"table_skip:{e}")

    # Comments via defusedxml
    try:
        with zipfile.ZipFile(io.BytesIO(blob)) as z:
            if "word/comments.xml" in z.namelist():
                root = et_fromstring(z.read("word/comments.xml"))
                for comment in root.iter(qn("w:comment")):
                    try:
                        author = comment.get(qn("w:author")) or "?"
                        texts = []
                        for t in comment.iter(qn("w:t")):
                            if t.text:
                                texts.append(t.text)
                        text = "".join(texts)
                        if text.strip():
                            sections.append(
                                Section(
                                    text=f"[{author}]: {text}",
                                    structure_path="Comment",
                                    section_type="comment",
                                    order=order,
                                )
                            )
                            order += 1
                    except Exception as e:
                        warnings.append(f"comment_xml_skip:{e}")
    except Exception as e:
        warnings.append(f"comments_zip_skip:{e}")

    for sec in doc.sections:
        try:
            for hdr_para in sec.header.paragraphs:
                if hdr_para.text.strip():
                    sections.append(
                        Section(text=hdr_para.text, structure_path="Header", section_type="header", order=order)
                    )
                    order += 1
            for ft_para in sec.footer.paragraphs:
                if ft_para.text.strip():
                    sections.append(
                        Section(text=ft_para.text, structure_path="Footer", section_type="footer", order=order)
                    )
                    order += 1
        except Exception as e:
            warnings.append(f"header_footer_skip:{e}")

    # Embedded xlsx in word/embeddings/
    try:
        from trimcp.extractors.office_excel import extract_xlsx_sync

        with zipfile.ZipFile(io.BytesIO(blob)) as z:
            for name in z.namelist():
                if not name.startswith("word/embeddings/"):
                    continue
                if not name.lower().endswith((".xlsx", ".xlsm")):
                    continue
                try:
                    data = z.read(name)
                    sub = extract_xlsx_sync(data)
                    if sub.text.strip():
                        sections.append(
                            Section(
                                text=f"[Embedded sheet: {name}]\n\n{sub.text}",
                                structure_path=f"Embedded: {name}",
                                section_type="body",
                                order=order,
                            )
                        )
                        order += 1
                    warnings.extend([f"embedded:{name}:{w}" for w in sub.warnings])
                except Exception as e:
                    warnings.append(f"embedded_xlsx_failed:{name}:{e}")
    except Exception as e:
        warnings.append(f"embedding_walk_skip:{e}")

    full_text = "\n\n".join(s.text for s in sections)
    if len(full_text.strip()) < 50:
        warnings.append("sparse_text:<50_chars; consider exporting with text layer or OCR pipeline (J.19)")

    metadata = _extract_core_props_xml(blob)
    full_text = "\n\n".join(s.text for s in sections)
    return ExtractionResult(
        method="python-docx",
        text=full_text,
        sections=sections,
        metadata=metadata,
        warnings=warnings,
    )


async def extract_doc(blob: bytes) -> ExtractionResult:
    converted = await asyncio.to_thread(libreoffice_convert, blob, ".doc", ".docx")
    if not converted:
        return empty_skipped("libreoffice", "conversion_failed", warnings=["doc conversion failed"])
    res = await extract_docx(converted)
    res.method = "libreoffice→python-docx"
    res.warnings.insert(0, "Converted from legacy .doc via LibreOffice")
    return res
